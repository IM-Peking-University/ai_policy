from __future__ import annotations

import asyncio
import json
import logging
import os
import re
import sys
import time
import traceback
import random
import hashlib
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Tuple
from contextlib import asynccontextmanager
from pathlib import Path

import tldextract
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse

import httpx
from pdfminer.high_level import extract_text as pdf_extract_text
from rapidfuzz import fuzz

from crawl4ai import AsyncWebCrawler
try:
    from crawl4ai.extraction_strategy import DefaultExtractionStrategy
except Exception:
    DefaultExtractionStrategy = None
from crawl4ai import CacheMode

# Relative path modification
RAW_ARCHIVE_DIR = Path("./raw_archive")
RAW_ARCHIVE_DIR.mkdir(parents=True, exist_ok=True)

DEFAULT_API_KEY = ""
DEFAULT_MODEL = "gpt-4o-mini"
INPUT_PATH = "./ai_policy_crawl_output_optimized.csv"
OUTPUT_PATH = "./1009_ai_policy_crawl_output_optimized.csv"

os.environ.setdefault("OPENAI_API_KEY", DEFAULT_API_KEY)
os.environ.setdefault("OPENAI_MODEL", DEFAULT_MODEL)

# Modified: Keep only English keywords
AI_KEYWORDS = [
    "ai", "artificial intelligence", "llm", "chatgpt", "gpt", "generative", "machine learning",
    "large language model", "generative ai", "chatbot", "transformer", "neural network"
]

# Modified: Keep only English keywords
POLICY_KEYWORDS = [
    "policy", "policies", "editorial", "ethics", "guideline", "guidelines",
    "instructions", "for authors", "submission", "authors", "publication ethics",
    "editorial policies", "author guidelines", "submission guidelines",
    "ethical guidelines", "publication policy", "manuscript preparation"
]

# Modified: Keep only English terms
AI_TERMS = [
    "ai", "artificial intelligence", "llm", "generative", "chatgpt", "gpt"
]

HUB_DOMAINS = [
    "springer.com", "wiley.com", "elsevier.com", "sagepub.com",
    "cambridge.org", "tandfonline.com", "nature.com", "sciencedirect.com",
    "oup.com", "biomedcentral.com", "ieee.org", "acm.org"
]

# Modified: Keep only English URL patterns
POLICY_URL_PATTERNS = [
    r"/author", r"/guideline", r"/policy", r"/instruction", r"/submission",
    r"/ethics", r"/editorial", r"/for-authors", r"/author-guidelines",
    r"/for[-_/]?authors", r"/instructions[-_/]?for[-_/]?authors", r"/submit",
    r"/guidelines?", r"/ethic"
]

NEGATIVE_AI_FALSES = [
    "aims & scope", "aims and scope", "ai ms", "art institute"
]

CROSS_DOMAIN_WHITELIST: Dict[str, List[str]] = {
    "wiley.com": ["onlinelibrary.wiley.com", "authorservices.wiley.com"],
    "tandfonline.com": ["tandfonline.com", "authorservices.taylorandfrancis.com"],
    "elsevier.com": ["elsevier.com", "www.elsevier.com", "www.editorialmanager.com"],
    "springer.com": ["springer.com", "link.springer.com", "springernature.com", "nature.com"],
    "sagepub.com": ["journals.sagepub.com", "mc.manuscriptcentral.com", "uk.sagepub.com"],
    "oup.com": ["academic.oup.com", "oup.com", "oxfordacademic.oup.com"],
    "ieee.org": ["ieee.org", "journals.ieeeauthorcenter.ieee.org"],
    "acm.org": ["dl.acm.org", "acm.org"],
}

def compute_keyword_stats(text: str) -> Tuple[List[str], int, bool]:
    lower_text = (text or "").lower()
    matched: List[str] = []
    total_count = 0
    
    # Exact match
    for kw in AI_KEYWORDS:
        pattern = re.escape(kw.lower()) if any(c.isalpha() for c in kw) else re.escape(kw)
        occurrences = len(re.findall(pattern, lower_text if any(c.isalpha() for c in kw) else (text or "")))
        if occurrences > 0:
            matched.append(kw)
            total_count += occurrences
    
    # Partial match
    for kw in AI_KEYWORDS:
        if len(kw) > 3 and kw not in matched:
            if kw.lower() in lower_text:
                matched.append(f"partial:{kw}")
                total_count += 1
    
    return matched, total_count, bool(matched)

def etld1(url: str) -> str:
    ext = tldextract.extract(url)
    return ".".join([p for p in [ext.domain, ext.suffix] if p])

def same_org(u1: str, u2: str) -> bool:
    return etld1(u1) == etld1(u2)

_DEF_UA = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
    "(KHTML, like Gecko) Chrome/118.0.0.0 Safari/537.36"
)

logger = logging.getLogger("ai_policy_crawler_optimized")
if not logger.handlers:
    handler = logging.StreamHandler(sys.stdout)
    formatter = logging.Formatter("[%(asctime)s] %(levelname)s: %(message)s")
    handler.setFormatter(formatter)
    logger.addHandler(handler)
logger.setLevel(logging.INFO)

ALLOWED_PAGE_TYPES = [
    "author_policies", "editorial_policies", "ethics", 
    "instructions_for_authors", "submissions", "publisher_hub", 
    "article", "misc"
]

@dataclass
class LinkCandidate:
    url: str
    anchor_text: str
    surrounding_text: str
    url_tokens: str
    priority: float = 0.0
    depth: int = 0

@dataclass
class ExtractedPolicy:
    policy_found: bool
    confidence: float
    policy_url: Optional[str]
    policy_text_excerpt: Optional[str]
    policy_summary: Optional[str]
    ai_usage_authoring: Optional[Dict[str, Any]]
    evidence_spans: List[Dict[str, Any]] = field(default_factory=list)

@dataclass
class OptimizedCrawlConfig:
    per_site_depth: int = 4
    per_site_max_pages: int = 25
    per_site_time_budget_s: int = 150
    concurrency: int = 4
    cache_mode: str = "ENABLED"
    allow_cross_domain_threshold: float = 0.65
    hit_threshold: float = 0.45
    frontier_topk: int = 40
    link_batch: int = 30
    llm_retry_attempts: int = 3
    llm_retry_delay: float = 2.0
    per_domain_concurrency: int = 2
    per_domain_delay_min: float = 0.2
    per_domain_delay_max: float = 0.6

class DomainRateLimiter:
    def __init__(self, cfg: OptimizedCrawlConfig):
        from collections import defaultdict
        self.cfg = cfg
        self._sems = defaultdict(lambda: asyncio.Semaphore(cfg.per_domain_concurrency))

    @asynccontextmanager
    async def slot(self, url: str):
        domain = etld1(url)
        sem = self._sems[domain]
        await sem.acquire()
        try:
            await asyncio.sleep(random.uniform(self.cfg.per_domain_delay_min, self.cfg.per_domain_delay_max))
            yield
        finally:
            sem.release()

class RawStorageManager:
    def __init__(self, base: Path = RAW_ARCHIVE_DIR):
        self.base = Path(base)
        self.base.mkdir(parents=True, exist_ok=True)

    def _hash(self, url: str) -> str:
        return hashlib.sha256(url.encode("utf-8")).hexdigest()[:16]

    def _dir_for(self, url: str) -> Path:
        d = self.base / etld1(url)
        d.mkdir(parents=True, exist_ok=True)
        return d

    def save_html(self, url: str, html: str) -> str:
        d = self._dir_for(url)
        fn = f"{self._hash(url)}.html"
        p = d / fn
        try:
            p.write_text(html or "", encoding="utf-8", errors="ignore")
            return str(p)
        except Exception as e:
            logger.warning(f"保存HTML失败 {url}: {e}")
            return ""

    def save_text(self, url: str, text: str) -> str:
        d = self._dir_for(url)
        fn = f"{self._hash(url)}.txt"
        p = d / fn
        try:
            p.write_text(text or "", encoding="utf-8", errors="ignore")
            return str(p)
        except Exception as e:
            logger.warning(f"保存TXT失败 {url}: {e}")
            return ""

    def save_binary(self, url: str, content: bytes, ext: str) -> str:
        d = self._dir_for(url)
        fn = f"{self._hash(url)}.{ext.lstrip('.')}"
        p = d / fn
        try:
            p.write_bytes(content or b"")
            return str(p)
        except Exception as e:
            logger.warning(f"保存二进制失败 {url}: {e}")
            return ""

class ResilientLLMClient:
    def __init__(self, model: Optional[str] = None, temperature: float = 0.0):
        from openai import OpenAI
        base_url = "https://api.openai.com/v1"
        self.client = OpenAI(base_url=base_url) if base_url else OpenAI()
        self.model = model or os.getenv("OPENAI_MODEL", "gpt-4o-mini")
        self.temperature = temperature

    async def acomplete_json(self, system_prompt: str, user_prompt: str, max_retries: int = 3) -> Dict[str, Any]:
        def _run():
            resp = self.client.chat.completions.create(
                model=self.model,
                temperature=self.temperature,
                response_format={"type": "json_object"},
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
            )
            return resp.choices[0].message.content

        for attempt in range(max_retries):
            try:
                content = await asyncio.to_thread(_run)
                return json.loads(content)
            except Exception as e:
                if attempt < max_retries - 1:
                    wait_time = 2 ** attempt
                    logger.warning(f"LLM调用失败，{wait_time}秒后重试 (尝试 {attempt + 1}/{max_retries}): {e}")
                    await asyncio.sleep(wait_time)
                    continue
                else:
                    logger.warning(f"LLM调用最终失败，使用启发式方法: {e}")
                    raise

async def fetch_pdf_text(url: str, timeout: int = 30) -> str:
    try:
        async with httpx.AsyncClient(timeout=timeout, headers={"User-Agent": _DEF_UA}) as client:
            r = await client.get(url)
            r.raise_for_status()
            content_type = r.headers.get("content-type", "").lower()
            if "application/pdf" not in content_type and not url.lower().endswith(".pdf"):
                return ""
            import tempfile, os
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(r.content)
                tmp_path = tmp.name
            try:
                txt = pdf_extract_text(tmp_path, laparams=None) or ""
                if not txt.strip():
                    txt = pdf_extract_text(tmp_path, laparams=None, codec='utf-8') or ""
            finally:
                os.unlink(tmp_path)
            return txt
    except Exception as e:
        logger.warning(f"PDF获取失败 {url}: {e}")
        return ""

async def fetch_pdf_text_and_save(url: str, saver: RawStorageManager, limiter: DomainRateLimiter, timeout: int = 30) -> Tuple[str, str]:
    saved_path = ""
    text = ""
    try:
        async with limiter.slot(url):
            async with httpx.AsyncClient(timeout=timeout, headers={"User-Agent": _DEF_UA}) as client:
                r = await client.get(url)
                r.raise_for_status()
                ctype = (r.headers.get("content-type") or "").lower()
                if ("pdf" not in ctype) and (not url.lower().endswith(".pdf")):
                    return "", ""
                saved_path = saver.save_binary(url, r.content, "pdf")
                import tempfile, os
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                    tmp.write(r.content)
                    tmp_path = tmp.name
                try:
                    text = pdf_extract_text(tmp_path, laparams=None) or ""
                    if not text.strip():
                        text = pdf_extract_text(tmp_path, laparams=None, codec='utf-8') or ""
                finally:
                    os.unlink(tmp_path)
    except Exception as e:
        logger.warning(f"PDF获取失败 {url}: {e}")
    return text, saved_path

async def fetch_docx_text(url: str, timeout: int = 30) -> str:
    try:
        async with httpx.AsyncClient(timeout=timeout, headers={"User-Agent": _DEF_UA}) as client:
            r = await client.get(url)
            r.raise_for_status()
            ctype = (r.headers.get("content-type") or "").lower()
            if not (url.lower().endswith(".docx") or "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in ctype or url.lower().endswith(".doc")):
                return ""
            import tempfile
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                tmp.write(r.content)
                p = tmp.name
            try:
                try:
                    from docx import Document
                except Exception:
                    return ""
                doc = Document(p)
                return "\n".join([para.text for para in doc.paragraphs])
            finally:
                os.unlink(p)
    except Exception as e:
        logger.warning(f"DOCX获取失败 {url}: {e}")
        return ""

async def fetch_docx_text_and_save(url: str, saver: RawStorageManager, limiter: DomainRateLimiter, timeout: int = 30) -> Tuple[str, str]:
    saved_path = ""
    text = ""
    try:
        async with limiter.slot(url):
            async with httpx.AsyncClient(timeout=timeout, headers={"User-Agent": _DEF_UA}) as client:
                r = await client.get(url)
                r.raise_for_status()
                ctype = (r.headers.get("content-type") or "").lower()
                if not (url.lower().endswith(".docx") or "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in ctype or url.lower().endswith(".doc")):
                    return "", ""
                saved_path = saver.save_binary(url, r.content, "docx")
                import tempfile
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                    tmp.write(r.content)
                    p = tmp.name
                try:
                    try:
                        from docx import Document
                    except Exception:
                        return "", saved_path
                    doc = Document(p)
                    text = "\n".join([para.text for para in doc.paragraphs])
                finally:
                    os.unlink(p)
    except Exception as e:
        logger.warning(f"DOCX获取失败 {url}: {e}")
    return text, saved_path

def extract_visible_text(html: str, max_chars: int = 60000) -> str:
    soup = BeautifulSoup(html, "lxml")
    for tag in soup(["script", "style", "noscript", "header", "footer", "nav"]):
        tag.decompose()
    text = soup.get_text(" ", strip=True)
    return text[:max_chars]

@dataclass
class PageParse:
    text: str
    links: List[LinkCandidate]
    alt_links: List[LinkCandidate] = field(default_factory=list)

def parse_page(html: str, base_url: str, limit_links: int = 250) -> PageParse:
    soup = BeautifulSoup(html, "lxml")
    links: List[LinkCandidate] = []

    for a in soup.find_all("a", href=True):
        href = a.get("href").strip()
        if href.startswith("#") or href.startswith("javascript:"):
            continue
        abs_url = urljoin(base_url, href)
        anchor = " ".join((a.get_text(" ") or "").split())[:300]

        parent = a.parent
        context = ""
        if parent:
            siblings_text = []
            for sibling in parent.previous_siblings:
                if hasattr(sibling, 'get_text'):
                    siblings_text.append(sibling.get_text(" ", strip=True))
            for sibling in parent.next_siblings:
                if hasattr(sibling, 'get_text'):
                    siblings_text.append(sibling.get_text(" ", strip=True))
            heading = parent.find_previous(["h1", "h2", "h3"])
            head_txt = heading.get_text(" ", strip=True) if heading else ""
            context = " ".join((siblings_text[-2:] + [head_txt, parent.get_text(" ", strip=True)]))

        context = context[:500]
        tokens = " ".join([etld1(abs_url), urlparse(abs_url).path.replace("/", " ")[:200]])
        links.append(LinkCandidate(url=abs_url, anchor_text=anchor, surrounding_text=context, url_tokens=tokens))
        if len(links) >= limit_links:
            break

    alt_links: List[LinkCandidate] = []
    for rel in soup.find_all("link", rel=True, href=True):
        rels = [r.lower() for r in (rel.get("rel") or [])]
        if "canonical" in rels:
            canu = urljoin(base_url, rel["href"])
            alt_links.append(LinkCandidate(url=canu, anchor_text="canonical", surrounding_text="", url_tokens=etld1(canu)))
    for suffix in ["?amp", "/amp", "?output=1", "/print"]:
        alt_links.append(LinkCandidate(url=base_url + suffix, anchor_text="alt", surrounding_text="", url_tokens=etld1(base_url)))

    text = extract_visible_text(html)
    return PageParse(text=text, links=links, alt_links=alt_links)

def extract_issn_like(text: str) -> List[str]:
    return list(set(re.findall(r"(?:ISSN|E-?ISSN)[:\s]*([0-9]{4}-[0-9Xx]{3}[0-9Xx])", text)))

def guess_publisher_templates(home_url: str, html_text: str) -> List[str]:
    urls = []
    host = etld1(home_url)
    path = urlparse(home_url).path.lower()
    issns = extract_issn_like(html_text)

    # Taylor & Francis: journalCode=
    m = re.search(r"journalCode=([a-z0-9]+)", html_text, re.I)
    jcode = m.group(1) if m else None
    if "tandfonline.com" in host:
        if jcode:
            urls.append(f"https://www.tandfonline.com/action/authorSubmission?journalCode={jcode}&show=instructions")
        urls.append(home_url.rstrip("/") + "/instructions-for-authors")

    # Wiley
    if "wiley.com" in host:
        for s in issns:
            s_clean = s.replace("-", "")
            urls.append(f"https://onlinelibrary.wiley.com/journal/{s_clean}/homepage/forauthors")
        urls.append(home_url.rstrip("/") + "/forauthors")

    # Elsevier
    if "elsevier.com" in host:
        last_seg = path.strip("/").split("/")[-1] if path.strip("/") else ""
        for s in issns:
            urls.append(f"https://www.elsevier.com/journals/{last_seg}/{s}/guide-for-authors")

    # Springer/Nature
    if "springer" in host or "nature.com" in host:
        m = re.search(r"/journal/(\d+)", home_url)
        if m:
            jid = m.group(1)
            urls += [
                f"https://www.springer.com/journal/{jid}/submissions",
                f"https://www.springer.com/journal/{jid}/guidelines",
            ]
        urls.append(home_url.rstrip("/") + "/authors")

    return [u for u in urls if u]

async def discover_from_sitemaps(home_url: str) -> List[str]:
    base = f"{urlparse(home_url).scheme}://{urlparse(home_url).netloc}"
    robots = f"{base}/robots.txt"
    out = []
    try:
        async with httpx.AsyncClient(headers={"User-Agent": _DEF_UA}, timeout=20) as cli:
            r = await cli.get(robots)
            if r.status_code != 200:
                return out
            smaps = re.findall(r"(?i)Sitemap:\s*(\S+)", r.text)
            smaps = smaps or [f"{base}/sitemap.xml", f"{base}/sitemap_index.xml"]
            for sm in smaps:
                try:
                    rs = await cli.get(sm)
                    if rs.status_code != 200:
                        continue
                    locs = re.findall(r"<loc>(.*?)</loc>", rs.text, re.I)
                    for loc in locs:
                        low = loc.lower()
                        # Modified: Keep only English terms
                        if any(p in low for p in ["author", "guideline", "submission", "ethic", "editorial", "for-authors"]):
                            out.append(loc.strip())
                except:
                    continue
    except:
        pass
    seen = set()
    result = []
    for u in out:
        if u not in seen:
            seen.add(u)
            result.append(u)
        if len(result) >= 80:
            break
    return result

def heuristic_score_link(lc: "LinkCandidate") -> Tuple[float, bool]:
    s = f"{(lc.anchor_text or '').lower()}\n{(lc.surrounding_text or '').lower()}\n{(lc.url_tokens or '').lower()}"
    prob = 0.0

    url_lower = (lc.url or '').lower()
    for pattern in POLICY_URL_PATTERNS:
        if re.search(pattern, url_lower):
            prob += 0.35
            break

    for kw in POLICY_KEYWORDS:
        if kw in s:
            prob += 0.15

    if any(kw in s for kw in AI_TERMS):
        prob += 0.2

    # Modified: Keep only English terms
    for kw in ("for authors", "instructions", "guideline", "ethic"):
        if fuzz.partial_ratio(kw, s) >= 85:
            prob += 0.1
            break

    if len(lc.anchor_text or '') > 20:
        prob += 0.05

    prob = max(0.0, min(1.0, prob))

    is_hub = any(d in url_lower for d in HUB_DOMAINS) or ('/journals' in url_lower or '/journal' in url_lower)

    return prob, is_hub

def heuristic_score_links(batch: List["LinkCandidate"]) -> List[Tuple[int, float, bool]]:
    out: List[Tuple[int, float, bool]] = []
    for i, lc in enumerate(batch):
        p, hub = heuristic_score_link(lc)
        out.append((i, p, hub))
    return out

def heuristic_extract_policy(text: str) -> Tuple[bool, float, str, str, List[str]]:
    if not text:
        return False, 0.0, "", "", []

    lower = text.lower()

    for neg in NEGATIVE_AI_FALSES:
        if neg in lower:
            lower = lower.replace(" ai ", " ")

    ai_terms_found = []
    for term in AI_TERMS:
        if term in lower:
            ai_terms_found.append(term)

    if not ai_terms_found:
        return False, 0.0, "", "", []

    policy_contexts = []
    sentences = re.split(r'[.!?。！？\n\r]+', text)

    for sentence in sentences:
        sentence_lower = sentence.lower().strip()
        if not sentence_lower:
            continue
        has_ai = any(term in sentence_lower for term in ai_terms_found)
        has_policy = any(kw in sentence_lower for kw in POLICY_KEYWORDS)
        if has_ai and has_policy:
            policy_contexts.append(sentence.strip())
        elif has_ai and len(sentence) > 30:
            policy_contexts.append(sentence.strip())

    if not policy_contexts:
        for sentence in sentences:
            sentence_lower = sentence.lower().strip()
            if any(term in sentence_lower for term in ai_terms_found):
                policy_contexts.append(sentence.strip())
            if len(policy_contexts) >= 5:
                break

    if not policy_contexts:
        return False, 0.0, "", "", []

    confidence = min(0.7, 0.3 + 0.1 * len(ai_terms_found) + 0.05 * len(policy_contexts))
    summary = policy_contexts[0][:250] + ("..." if len(policy_contexts[0]) > 250 else "")
    excerpt = " | ".join(policy_contexts[:3])[:1000]
    return True, confidence, summary, excerpt, policy_contexts[:5]

def build_extraction_prompt(url: str, page_text: str) -> Tuple[str, str]:
    system = "You extract structured policy facts about authors' use of AI/LLMs in writing. Return ONLY JSON."
    user = f"""
Extract the policy regarding authors' use of AI/LLMs in writing from the page.
Look for any mention of AI, artificial intelligence, large language models, ChatGPT, etc.
in the context of author guidelines, submission policies, or editorial policies.

URL: {url}
PAGE TEXT (truncated):
{page_text[:18000]}

Return JSON:
{{
  "policy_found": boolean,
  "confidence": 0..1,
  "policy_summary": "one sentence summary if found else empty",
  "policy_text_excerpt": "short key excerpt if found",
  "ai_usage_authoring": {{
     "category": "allow_with_disclosure|allow_limited|prohibit|unclear",
     "disclosure": "required|recommended|not_required|unspecified",
     "exceptions": "string or empty"
  }},
  "evidence_quotes": ["up to 3 short quotes that justify the extraction"]
}}
Only JSON.
"""
    return system, user

def build_page_classify_prompt(url: str, page_text: str) -> Tuple[str, str]:
    system = "You are a precise classifier for scholarly journal websites. Return ONLY JSON."
    user = f"""
Task: Classify the page and estimate whether it likely contains policy text about
*authors' use of AI/LLMs in writing*. Be more lenient in identifying potential policy pages.

URL: {url}

Page excerpt:
{page_text[:15000]}

Respond as JSON with:
{{
  "page_type": "one of {ALLOWED_PAGE_TYPES}",
  "ai_policy_presence_prob": 0..1,
  "publisher_hub_prob": 0..1,
  "short_summary": "<=280 chars summary of what this page is about",
  "reasoning": "brief reasoning for classification"
}}
Only JSON.
"""
    return system, user

def allow_cross_domain(home: str, link: str) -> bool:
    h1, h2 = etld1(home), etld1(link)
    if h1 == h2:
        return True
    for base, allowed in CROSS_DOMAIN_WHITELIST.items():
        if base in h1:
            for d in allowed:
                if etld1(link).endswith(etld1(d)) or d in link:
                    return True
    return False

class OptimizedAIPolicyCrawler:
    def __init__(self, llm: ResilientLLMClient, config: OptimizedCrawlConfig):
        self.llm = llm
        self.cfg = config
        self.llm_disabled: bool = False
        self.llm_failures: int = 0
        self.max_llm_failures: int = 10
        self.rate_limiter = DomainRateLimiter(config)
        self.saver = RawStorageManager(RAW_ARCHIVE_DIR)

    async def _arender(self, crawler: AsyncWebCrawler, url: str) -> Tuple[str, str]:
        mode_str = (self.cfg.cache_mode or "ENABLED").upper()
        cm_default = getattr(CacheMode, "DEFAULT", None) or getattr(CacheMode, "ENABLED", None)
        cm_bypass = getattr(CacheMode, "BYPASS", None) or getattr(CacheMode, "DISABLED", None)
        cm_only = getattr(CacheMode, "CACHE_ONLY", None) or getattr(CacheMode, "ONLY", None)
        cache_mode = None
        if mode_str == "ENABLED":
            cache_mode = cm_default
        elif mode_str == "BYPASS":
            cache_mode = cm_bypass
        elif mode_str == "ONLY":
            cache_mode = cm_only
        else:
            cache_mode = cm_default
        _kwargs = {"url": url}
        if cache_mode is not None:
            _kwargs["cache_mode"] = cache_mode
        if mode_str == "BYPASS":
            _kwargs["bypass_cache"] = True
        if DefaultExtractionStrategy is not None:
            try:
                _kwargs["extraction_strategy"] = DefaultExtractionStrategy()
            except Exception:
                pass
        try:
            async with self.rate_limiter.slot(url):
                res = await crawler.arun(**_kwargs)
        except TypeError:
            async with self.rate_limiter.slot(url):
                res = await crawler.arun(url=url)
        html = getattr(res, "html", None) or getattr(res, "content", "") or ""
        md = getattr(res, "markdown", None) or getattr(res, "fit_markdown", "") or ""
        return html, md

    async def _classify_page(self, url: str, text: str) -> Dict[str, Any]:
        if self.llm_disabled or self.llm_failures >= self.max_llm_failures:
            return self._heuristic_classify_page(url, text)
        try:
            system, user = build_page_classify_prompt(url, text)
            result = await self.llm.acomplete_json(system, user, max_retries=self.cfg.llm_retry_attempts)
            self.llm_failures = 0
            return result
        except Exception as e:
            self.llm_failures += 1
            logger.warning(f"LLM分类失败 ({self.llm_failures}/{self.max_llm_failures}): {e}")
            if self.llm_failures >= self.max_llm_failures:
                self.llm_disabled = True
                logger.warning("LLM禁用，切换到启发式方法")
            return self._heuristic_classify_page(url, text)

    def _heuristic_classify_page(self, url: str, text: str) -> Dict[str, Any]:
        lower_text = (text or "").lower()
        page_type = "misc"
        ai_policy_prob = 0.0
        hub_prob = 0.0

        url_lower = url.lower()
        if any(pattern in url_lower for pattern in POLICY_URL_PATTERNS):
            page_type = "author_policies"
            ai_policy_prob = 0.5

        ai_terms_count = sum(1 for term in AI_TERMS if term in lower_text)
        policy_terms_count = sum(1 for term in POLICY_KEYWORDS if term in lower_text)

        if ai_terms_count > 0 and policy_terms_count > 0:
            page_type = "editorial_policies"
            ai_policy_prob = min(0.85, 0.3 + 0.1 * ai_terms_count + 0.1 * policy_terms_count)
        elif ai_terms_count > 0:
            ai_policy_prob = min(0.6, 0.2 + 0.1 * ai_terms_count)

        if any(domain in url_lower for domain in HUB_DOMAINS) or "/journals" in url_lower:
            hub_prob = 0.8

        summary = f"Heuristic: {page_type}, AI terms: {ai_terms_count}, Policy terms: {policy_terms_count}"
        return {
            "page_type": page_type,
            "ai_policy_presence_prob": ai_policy_prob,
            "publisher_hub_prob": hub_prob,
            "short_summary": summary,
            "reasoning": "Heuristic classification based on keywords and URL patterns"
        }

    async def _score_links(self, url: str, links: List["LinkCandidate"]) -> List[Tuple[int, float, bool]]:
        return heuristic_score_links(links)

    async def _extract_policy(self, url: str, text: str) -> "ExtractedPolicy":
        found, conf, summary, excerpt, quotes = heuristic_extract_policy(text)
        if found and conf >= self.cfg.hit_threshold:
            spans = []
            for q in quotes:
                if not q:
                    continue
                idx = text.find(q)
                if idx == -1:
                    idx = _approx_find(text, q)
                if idx != -1:
                    spans.append({"quote": q, "start": idx, "end": idx + len(q), "page_url": url})
            return ExtractedPolicy(
                policy_found=True,
                confidence=conf,
                policy_url=url,
                policy_text_excerpt=excerpt or None,
                policy_summary=summary or None,
                ai_usage_authoring={
                    "category": "detected_by_heuristic",
                    "disclosure": "unspecified",
                    "exceptions": ""
                },
                evidence_spans=spans,
            )
        
        if not self.llm_disabled and self.llm_failures < self.max_llm_failures:
            try:
                system, user = build_extraction_prompt(url, text)
                data = await self.llm.acomplete_json(system, user, max_retries=self.cfg.llm_retry_attempts)
                quotes = data.get("evidence_quotes", []) or []
                spans = []
                for q in quotes:
                    if not q:
                        continue
                    idx = text.find(q)
                    if idx == -1:
                        idx = _approx_find(text, q)
                    if idx != -1:
                        spans.append({"quote": q, "start": idx, "end": idx + len(q), "page_url": url})
                return ExtractedPolicy(
                    policy_found=bool(data.get("policy_found", False)),
                    confidence=float(data.get("confidence", 0.0)),
                    policy_url=url,
                    policy_text_excerpt=data.get("policy_text_excerpt"),
                    policy_summary=data.get("policy_summary"),
                    ai_usage_authoring=data.get("ai_usage_authoring"),
                    evidence_spans=spans,
                )
            except Exception as e:
                self.llm_failures += 1
                logger.warning(f"LLM提取失败: {e}")

        return ExtractedPolicy(
            policy_found=False,
            confidence=0.0,
            policy_url=None,
            policy_text_excerpt=None,
            policy_summary=None,
            ai_usage_authoring=None,
            evidence_spans=[],
        )

    def _policy_like(self, page_type: str) -> bool:
        return page_type in {
            "author_policies", "editorial_policies", "ethics", "instructions_for_authors", "submissions"
        }

    async def crawl_one(self, homepage_url: str) -> Dict[str, Any]:
        start_ts = time.time()
        per_site_pages = 0
        visited: set[str] = set()
        frontier: List[Tuple[float, int, LinkCandidate]] = []
        tiebreak = 0

        result_record: Dict[str, Any] = {
            "homepage_url": homepage_url,
            "policy_found": False,
            "policy_url": None,
            "policy_summary": None,
            "policy_text_excerpt": None,
            "ai_usage_authoring": None,
            "evidence_spans": [],
            "confidence": 0.0,
            "elapsed_s": None,
            "visited_urls": [],
            "raw_saved_paths": [],
            "visited_pages": 0,
        }

        org = etld1(homepage_url)
        logger.info(f"开始爬取: {org}")

        async with AsyncWebCrawler() as crawler:
            parse = None
            try:
                html, md = await self._arender(crawler, homepage_url)
                per_site_pages += 1
                if html:
                    p_html = self.saver.save_html(homepage_url, html)
                    if p_html:
                        result_record["raw_saved_paths"].append(p_html)
                if md:
                    p_txt = self.saver.save_text(homepage_url, md)
                    if p_txt:
                        result_record["raw_saved_paths"].append(p_txt)

                parse = parse_page(html or md, homepage_url)
                cls = await self._classify_page(homepage_url, parse.text)
                if self._policy_like(cls.get("page_type", "misc")):
                    if float(cls.get("ai_policy_presence_prob", 0.0)) >= self.cfg.hit_threshold:
                        pol = await self._extract_policy(homepage_url, parse.text)
                        if pol.policy_found:
                            _fill_result(result_record, pol)
                            result_record["policy_url"] = homepage_url
                            result_record["elapsed_s"] = round(time.time() - start_ts, 3)
                            result_record["visited_pages"] = per_site_pages
                            result_record["visited_urls"].append(homepage_url)
                            return result_record
            except Exception as e:
                logger.warning(f"Failue {homepage_url}: {e}")

            try:
                sm_urls = await discover_from_sitemaps(homepage_url)
                for u in sm_urls:
                    _frontier_push(frontier, LinkCandidate(u, "sitemap", "", etld1(u)), 0.95, tiebreak, self.cfg.frontier_topk)
                    tiebreak += 1
            except Exception as e:
                logger.debug(f"Failue: {e}")

            try:
                if parse:
                    tmpl = guess_publisher_templates(homepage_url, (html or md or ""))
                    for u in tmpl:
                        _frontier_push(frontier, LinkCandidate(u, "publisher-template", "", etld1(u)), 0.9, tiebreak, self.cfg.frontier_topk)
                        tiebreak += 1
                    for alt in (parse.alt_links or []):
                        _frontier_push(frontier, alt, 0.85, tiebreak, self.cfg.frontier_topk)
                        tiebreak += 1
            except Exception as e:
                logger.debug(f"Failue: {e}")

            if parse:
                scored = await self._score_links(homepage_url, parse.links)
                for (idx, prob, is_pub) in scored:
                    lc = parse.links[idx]
                    allow = same_org(homepage_url, lc.url) or allow_cross_domain(homepage_url, lc.url) or (is_pub and prob >= self.cfg.allow_cross_domain_threshold)
                    if not allow:
                        continue
                    lc.priority = prob
                    lc.depth = 1
                    _frontier_push(frontier, lc, prob, tiebreak, self.cfg.frontier_topk)
                    tiebreak += 1

            while frontier:
                if per_site_pages >= self.cfg.per_site_max_pages:
                    logger.info(f"Limiti: {org}")
                    break
                if time.time() - start_ts > self.cfg.per_site_time_budget_s:
                    logger.info(f"Limit: {org}")
                    break

                prio, _, node = frontier.pop(0)
                url = node.url
                if url in visited:
                    continue
                if node.depth > self.cfg.per_site_depth:
                    continue
                visited.add(url)

                try:
                    # PDF
                    if url.lower().endswith(".pdf"):
                        text, savep = await fetch_pdf_text_and_save(url, self.saver, self.rate_limiter)
                        if savep:
                            result_record["raw_saved_paths"].append(savep)
                        if text:
                            cls = await self._classify_page(url, text)
                            if float(cls.get("ai_policy_presence_prob", 0.0)) >= self.cfg.hit_threshold:
                                pol = await self._extract_policy(url, text)
                                if pol.policy_found:
                                    _fill_result(result_record, pol)
                                    result_record["policy_url"] = url
                                    result_record["elapsed_s"] = round(time.time() - start_ts, 3)
                                    result_record["visited_pages"] = per_site_pages
                                    result_record["visited_urls"].append(url)
                                    return result_record
                        result_record["visited_urls"].append(url)
                        continue

                    # DOCX
                    if url.lower().endswith(".docx") or url.lower().endswith(".doc"):
                        text, savep = await fetch_docx_text_and_save(url, self.saver, self.rate_limiter)
                        if savep:
                            result_record["raw_saved_paths"].append(savep)
                        if text:
                            cls = await self._classify_page(url, text)
                            if self._policy_like(cls.get("page_type", "misc")) and float(cls.get("ai_policy_presence_prob", 0.0)) >= self.cfg.hit_threshold:
                                pol = await self._extract_policy(url, text)
                                if pol.policy_found:
                                    _fill_result(result_record, pol)
                                    result_record["policy_url"] = url
                                    result_record["elapsed_s"] = round(time.time() - start_ts, 3)
                                    result_record["visited_pages"] = per_site_pages
                                    result_record["visited_urls"].append(url)
                                    return result_record
                        result_record["visited_urls"].append(url)
                        continue

                    html, md = await self._arender(crawler, url)
                    per_site_pages += 1
                    if html:
                        p_html = self.saver.save_html(url, html)
                        if p_html:
                            result_record["raw_saved_paths"].append(p_html)
                    if md:
                        p_txt = self.saver.save_text(url, md)
                        if p_txt:
                            result_record["raw_saved_paths"].append(p_txt)

                    parse2 = parse_page(html or md, url)
                    cls = await self._classify_page(url, parse2.text)

                    if self._policy_like(cls.get("page_type", "misc")):
                        if float(cls.get("ai_policy_presence_prob", 0.0)) >= self.cfg.hit_threshold:
                            pol = await self._extract_policy(url, parse2.text)
                            if pol.policy_found:
                                _fill_result(result_record, pol)
                                result_record["policy_url"] = url
                                result_record["elapsed_s"] = round(time.time() - start_ts, 3)
                                result_record["visited_pages"] = per_site_pages
                                result_record["visited_urls"].append(url)
                                return result_record

                    for alt in (parse2.alt_links or []):
                        _frontier_push(frontier, alt, 0.85, tiebreak, self.cfg.frontier_topk)
                        tiebreak += 1

                    scored = await self._score_links(url, parse2.links)
                    for (idx, prob, is_pub) in scored:
                        lc = parse2.links[idx]
                        allow = same_org(homepage_url, lc.url) or allow_cross_domain(homepage_url, lc.url) or (is_pub and prob >= self.cfg.allow_cross_domain_threshold)
                        if not allow:
                            continue
                        new_depth = node.depth + 1
                        if new_depth > self.cfg.per_site_depth:
                            continue
                        lc.priority = prob
                        lc.depth = new_depth
                        _frontier_push(frontier, lc, prob, tiebreak, self.cfg.frontier_topk)
                        tiebreak += 1

                    result_record["visited_urls"].append(url)

                except Exception as e:
                    logger.warning(f"Failue {url}: {e}")
                    continue

        result_record["elapsed_s"] = round(time.time() - start_ts, 3)
        result_record["visited_pages"] = per_site_pages
        return result_record

def _approx_find(text: str, quote: str) -> int:
    q = quote.strip()
    if not q:
        return -1
    best_idx, best_score = -1, 0
    window = max(20, min(200, len(q) + 20))
    for i in range(0, max(1, len(text) - window), max(5, window // 2)):
        seg = text[i : i + window]
        score = fuzz.partial_ratio(q, seg)
        if score > best_score:
            if q.split():
                first_token = q.split()[0]
                local = seg.find(first_token)
                offset = local if local >= 0 else 0
            else:
                offset = 0
            best_score, best_idx = score, i + offset
    return best_idx if best_score >= 80 else -1

def _frontier_push(frontier: List[Tuple[float, int, LinkCandidate]], lc: LinkCandidate, prob: float, tiebreak: int, topk: Optional[int] = None):
    frontier.append((-prob, tiebreak, lc))
    frontier.sort(key=lambda x: (x[0], x[1]))
    if isinstance(topk, int) and topk > 0 and len(frontier) > topk:
        del frontier[topk:]

def _fill_result(dst: Dict[str, Any], pol: ExtractedPolicy):
    dst.update({
        "policy_found": pol.policy_found,
        "confidence": pol.confidence,
        "policy_summary": pol.policy_summary,
        "policy_text_excerpt": pol.policy_text_excerpt,
        "ai_usage_authoring": pol.ai_usage_authoring,
        "evidence_spans": pol.evidence_spans,
    })

import csv

REQUIRED_FIELDS = [
    "journal_name", "url", "status", "http_status", "word_count",
    "has_ai_keywords", "matched_keywords", "keyword_match_count", 
    "ai_policy_text", "error", "followed_links"
]

async def run_optimized_crawl():
    cfg = OptimizedCrawlConfig()
    llm = ResilientLLMClient(model=DEFAULT_MODEL, temperature=0.0)
    crawler = OptimizedAIPolicyCrawler(llm, cfg)

    processed_set: set[str] = set()
    header_written = False
    
    if os.path.exists(OUTPUT_PATH) and os.path.getsize(OUTPUT_PATH) > 0:
        try:
            with open(OUTPUT_PATH, newline="", encoding="utf-8") as f:
                reader = csv.DictReader(f)
                for row in reader:
                    url = (row.get("url") or "").strip()
                    if url:
                        processed_set.add(url)
            header_written = True
            logger.info(f"process: {len(processed_set)}")
        except Exception as e:
            logger.warning(f"Failue: {e}")

    if not header_written:
        with open(OUTPUT_PATH, "w", newline="", encoding="utf-8") as f:
            writer = csv.DictWriter(f, fieldnames=REQUIRED_FIELDS)
            writer.writeheader()

    write_lock = asyncio.Lock()
    sem = asyncio.Semaphore(max(1, cfg.concurrency))

    async def append_row_safe(row: Dict[str, Any]):
        async with write_lock:
            with open(OUTPUT_PATH, "a", newline="", encoding="utf-8") as f:
                writer = csv.DictWriter(f, fieldnames=REQUIRED_FIELDS)
                writer.writerow({k: row.get(k, "") for k in REQUIRED_FIELDS})

    async def process_row(idx: int, row: Dict[str, Any]):
        async with sem:
            for col in REQUIRED_FIELDS:
                if col not in row:
                    row[col] = ""

            input_url = (row.get("url") or "").strip()
            if not input_url:
                logger.info(f"[{idx}/?]")
                if input_url not in processed_set:
                    await append_row_safe(row)
                return

            if input_url in processed_set:
                logger.info(f"[{idx}/?] | {input_url}")
                return

            if (row.get("ai_policy_text") or "").strip():
                logger.info(f"[{idx}/?] | {input_url}")
                await append_row_safe(row)
                processed_set.add(input_url)
                return

            start_ts = time.time()
            
            try:
                rec = await asyncio.wait_for(crawler.crawl_one(input_url), timeout=240)

                matched, match_cnt, has_kw = compute_keyword_stats(rec.get("policy_summary") or rec.get("policy_text_excerpt") or "")
                row["has_ai_keywords"] = str(bool(has_kw))
                row["matched_keywords"] = ",".join(matched)
                row["keyword_match_count"] = str(match_cnt)
                row["followed_links"] = json.dumps({
                    "visited": rec.get("visited_urls", []),
                    "raw_saved": rec.get("raw_saved_paths", []),
                    "visited_pages": rec.get("visited_pages", 0)
                }, ensure_ascii=False)

                new_text = rec.get("policy_summary") or rec.get("policy_text_excerpt") or ""

                await append_row_safe(row)
                processed_set.add(input_url)
            except asyncio.TimeoutError:
                row["error"] = (row.get("error") or "") + (" | " if row.get("error") else "") + "timeout"
                await append_row_safe(row)
                processed_set.add(input_url)
                logger.warning(f"[{idx}/?] {input_url}")
            except Exception as e:
                row["error"] = (row.get("error") or "") + (" | " if row.get("error") else "") + f"exception: {str(e)}"
                await append_row_safe(row)
                processed_set.add(input_url)
                logger.error(f"[{idx}/?] | {input_url}\n{traceback.format_exc()}")

    tasks = []
    with open(INPUT_PATH, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        rows = list(reader)
        total_rows = len(rows)
        for idx, row in enumerate(rows, 1):
            tasks.append(asyncio.create_task(process_row(idx, dict(row))))
            if len(tasks) >= cfg.concurrency * 2:
                await asyncio.gather(*tasks)
                tasks.clear()
            if idx % 10 == 0:
                logger.info(f"{idx}/{total_rows}")
    if tasks:
        await asyncio.gather(*tasks)

    logger.info(f"{total_rows} | {len(processed_set)}")

if __name__ == "__main__":
    asyncio.run(run_optimized_crawl())